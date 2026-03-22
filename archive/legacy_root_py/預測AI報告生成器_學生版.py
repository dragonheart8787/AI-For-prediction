#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SuperFusionAGI 預測AI系統報告生成器（學生版）
口吻目標：
- 以「大學生提交課堂報告」為語氣，避免推銷與過度確定
- 使用第一人稱敘述（我在實作、我嘗試、我觀察到）
- 主動標示假設、限制、未完成與需要改進的部分
- 避免表情符號與誇張形容詞
輸出：SuperFusionAGI_課堂報告_劉哲廷.docx
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn


def add_heading_center(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def build_student_report(output_name: str = 'SuperFusionAGI_課堂報告_劉哲廷.docx') -> str:
    doc = Document()

    # Base font for Chinese
    style = doc.styles['Normal']
    style.font.name = 'PMingLiU'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
    style.font.size = Pt(12)

    # 封面
    add_heading_center(doc, 'SuperFusionAGI：預測系統期末報告', 0)
    add_heading_center(doc, '課程：人工智慧實務（2024-2025學年）', 1)
    p = doc.add_paragraph('學生：劉哲廷')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)
    add_paragraph(doc, '')

    # 摘要（學生語氣）
    add_heading_center(doc, '摘要', 1)
    add_paragraph(doc, (
        '這份報告記錄我在學期內嘗試實作一個「多模型融合的預測系統」的過程。整體上，我先整合了' 
        '常見的表格模型（XGBoost、LightGBM）與簡單的序列模型，之後再把推理流程做成統一介面。' 
        '在訓練與測試過程中，我盡量保持條件可重複，並把失敗與未解決的地方一併寫進來。' 
        '最終模型在我使用的資料上有一定表現，但我仍不確定是否適用於更廣泛的場景，需要進一步驗證。'
    ))

    # 目錄（簡化，讓老師快速定位）
    add_heading_center(doc, '目錄', 1)
    for item in [
        '1. 研究動機與目標',
        '2. 系統架構（我怎麼把元件湊起來）',
        '3. 實作細節與假設',
        '4. 實驗設定與結果（含失敗案例）',
        '5. 限制與風險',
        '6. 我學到的事（反思）',
    ]:
        doc.add_paragraph(item, style='List Number')

    # 1. 動機與目標
    doc.add_heading('1. 研究動機與目標', 1)
    add_paragraph(doc, (
        '起初我想做一個「不挑資料源」的預測介面，原因是我常在專題中面臨資料格式不一致與模型切換麻煩。' 
        '因此我設定的目標不是要比所有方法更準，而是把「整合、可重複、可維護」放在第一位。' 
        '我也嘗試讓系統能在沒有 GPU 的情況下仍可執行，這就是我選擇 ONNX 的主要理由之一。'
    ))

    # 2. 架構
    doc.add_heading('2. 系統架構（我怎麼把元件湊起來）', 1)
    add_paragraph(doc, (
        '系統分成三層：資料取得與處理、模型訓練與融合、推理與服務。資料層我先用幾個常見來源做假資料流程，' 
        '模型層則以 XGBoost 與 LightGBM 為主，序列資料用簡化的 LSTM；推理層我做了一個 UnifiedPredictor，把輸入、' 
        '特徵處理與批量推理綁成固定規格。這樣做的缺點是靈活度會下降，但在作業與專題情境比較容易複製與除錯。'
    ))

    # 3. 實作細節與假設
    doc.add_heading('3. 實作細節與假設', 1)
    add_paragraph(doc, (
        '（1）資料假設：我預設輸入至少包含時間戳與數個連續特徵，若缺欄位我會先做簡單補值；' 
        '（2）模型設定：樹模型採用預設超參數做基準，僅在少數情況做網格或貝葉斯搜尋；' 
        '（3）ONNX：我先以 CPU 推理為主，嘗試 batch 推理（通常 512~2048）以換取吞吐量；' 
        '（4）評測：以切分的驗證集與簡單回測為主。我知道這樣會忽略不少真實世界的雜訊，但在學期時間內我先把流程做順。'
    ))

    # 4. 實驗與結果（保留保守語氣）
    doc.add_heading('4. 實驗設定與結果（含失敗案例）', 1)
    add_paragraph(doc, (
        '在幾個示範資料集上，我把單一模型與融合結果做了對比。整體趨勢看起來，融合在穩定度上略有優勢，' 
        '但在某些特徵工程不足的資料上，樹模型單獨表現並沒有比較差。這讓我意識到「資料品質」在這個題目上比模型更關鍵。'
    ))
    add_paragraph(doc, (
        '失敗部分主要出現在：序列長度設置不當（導致 LSTM 過擬合）、特徵標準化順序錯誤、以及 ONNX 匯出後的數值' 
        '偏差。我最後選擇優先保留可重複的樹模型流程，序列模型只在需要時打開。這樣雖不漂亮，但至少能穩定交作業。'
    ))

    # 5. 限制與風險（明確承認）
    doc.add_heading('5. 限制與風險', 1)
    doc.add_paragraph('以下是我目前明確知道的限制：')
    for pt in [
        '資料來源單一：本學期沒有拿到大型真實資料，外推性存疑。',
        '特徵工程簡化：部分特徵處理為了趕進度採用保守做法，可能壓低上限。',
        '序列模型未充分調參：時間有限，沒有做完整的超參數尋找。',
        'ONNX 精確度：個別模型在匯出後可能與原生推理有 ±(1e-4~1e-3) 的差距。',
        '評測設定偏理想：回測切分與資料清理相對乾淨，與真實場景仍有落差。',
    ]:
        doc.add_paragraph(pt, style='List Bullet')

    # 6. 反思（學生式）
    doc.add_heading('6. 我學到的事（反思）', 1)
    add_paragraph(doc, (
        '這次實作最大的收穫，是認識到「把東西做穩」比「堆方法」更難。我一開始想一次把很多模型都加進來，' 
        '最後發現只要資料一變，維護成本就會急劇上升。後來我調整做法：先把統一介面和基本評測做好，' 
        '再慢慢加模型與優化。這樣雖然進度看起來慢，但我比較能掌握每一步的影響。'
    ))
    add_paragraph(doc, (
        '如果下學期繼續做，我希望能：（1）找一個較完整的公開資料集做長期回測；（2）把特徵工程模組化；' 
        '（3）把 ONNX 與批量推理的誤差範圍再縮小；（4）做一份更嚴格的實驗記錄與隨機種子控制。'
    ))

    # 致謝（學生給教授）
    doc.add_heading('致謝', 1)
    add_paragraph(doc, (
        '感謝授課老師在課堂上的指導與提醒。本報告若有疏漏與錯誤，皆由我負責；' 
        '後續我會依據批改意見再修正程式與論述。'
    ))

    # 保存
    doc.save(output_name)
    return output_name


if __name__ == '__main__':
    out = build_student_report()
    print(f'Generated: {out}')

