#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整的量子廣義相對論 Word 文檔生成器
包含圖片插入和數學符號轉換
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.enum.section import WD_SECTION
    print("python-docx 已安裝")
except ImportError:
    print("安裝 python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.enum.section import WD_SECTION

def create_images_directory():
    """創建圖片目錄並下載圖片"""
    images_dir = Path("images")
    images_dir.mkdir(exist_ok=True)
    
    # 創建圖片說明文件
    image_descriptions = {
        "figure_2_curvature.png": "曲率分布圖 - 顯示熵場導出的時空彎曲結構",
        "figure_3_energy.png": "能量密度分布 - 幾何與能量守恆關係",
        "figure_4_convergence.png": "損失函數收斂曲線 - 模型達到協變與能量平衡",
        "figure_5_mapping.png": "模型層與理論層對映圖 - AI層與物理層的結構映射",
        "figure_6_quantum_correction.png": "量子修正項分布 - 熵梯度導致的局域量子修正",
        "figure_7_observer_field.png": "觀測者場效應 - 觀測行為對時空的回饋作用",
        "figure_8_hierarchy.png": "理論層次結構圖 - 從QIGG到CSC的遞階演化",
        "figure_9_equations.png": "統一方程組 - 最終的自洽理論系統",
        "figure_10_energy_table.png": "能量密度對比表 - 各理論層的能量平衡狀態",
        "figure_11_theory_comparison.png": "理論比較表 - 與現有理論的對比分析"
    }
    
    # 創建圖片描述文件
    with open(images_dir / "image_descriptions.txt", "w", encoding="utf-8") as f:
        f.write("圖片說明文件\n")
        f.write("=" * 50 + "\n\n")
        for filename, description in image_descriptions.items():
            f.write(f"{filename}: {description}\n")
    
    print(f"圖片目錄已創建：{images_dir}")
    print("請將對應的圖片文件放入此目錄")
    return images_dir

def convert_math_symbols(text):
    """轉換數學符號為可讀格式"""
    # 常見數學符號轉換
    math_conversions = {
        # 希臘字母
        r'\\mu': 'μ',
        r'\\nu': 'ν', 
        r'\\rho': 'ρ',
        r'\\lambda': 'λ',
        r'\\kappa': 'κ',
        r'\\Omega': 'Ω',
        r'\\nabla': '∇',
        r'\\Box': '□',
        r'\\Psi': 'Ψ',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\gamma': 'γ',
        r'\\delta': 'δ',
        r'\\epsilon': 'ε',
        r'\\zeta': 'ζ',
        r'\\eta': 'η',
        r'\\theta': 'θ',
        r'\\iota': 'ι',
        r'\\xi': 'ξ',
        r'\\pi': 'π',
        r'\\sigma': 'σ',
        r'\\tau': 'τ',
        r'\\phi': 'φ',
        r'\\chi': 'χ',
        r'\\psi': 'ψ',
        r'\\omega': 'ω',
        
        # 數學運算符
        r'\\partial': '∂',
        r'\\sum': '∑',
        r'\\int': '∫',
        r'\\infty': '∞',
        r'\\times': '×',
        r'\\cdot': '·',
        r'\\pm': '±',
        r'\\mp': '∓',
        r'\\leq': '≤',
        r'\\geq': '≥',
        r'\\neq': '≠',
        r'\\approx': '≈',
        r'\\equiv': '≡',
        r'\\propto': '∝',
        r'\\in': '∈',
        r'\\subset': '⊂',
        r'\\supset': '⊃',
        r'\\cup': '∪',
        r'\\cap': '∩',
        r'\\emptyset': '∅',
        r'\\rightarrow': '→',
        r'\\leftarrow': '←',
        r'\\leftrightarrow': '↔',
        r'\\Rightarrow': '⇒',
        r'\\Leftarrow': '⇐',
        r'\\Leftrightarrow': '⇔',
        
        # 括號
        r'\\left\\(': '(',
        r'\\right\\)': ')',
        r'\\left\\[': '[',
        r'\\right\\]': ']',
        r'\\left\\{': '{',
        r'\\right\\}': '}',
        
        # 分數
        r'\\frac\{([^}]+)\}\{([^}]+)\}': r'\1/\2',
        
        # 上下標簡化
        r'\^([0-9])': r'^\1',
        r'_([a-zA-Z0-9]+)': r'_\1',
        
        # 其他
        r'\\langle': '⟨',
        r'\\rangle': '⟩',
        r'\\mid': '|',
        r'\\cdot': '·',
        r'\\cdots': '⋯',
        r'\\ldots': '…',
        r'\\vdots': '⋮',
        r'\\ddots': '⋱'
    }
    
    for pattern, replacement in math_conversions.items():
        text = re.sub(pattern, replacement, text)
    
    return text

def add_image_to_document(doc, image_path, caption, width=4.5):
    """在文檔中添加圖片"""
    if os.path.exists(image_path):
        try:
            # 添加圖片
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(image_path, width=Inches(width))
            
            # 添加圖片標題
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(f"圖 {caption}")
            caption_run.font.size = Pt(11)
            caption_run.italic = True
            
            doc.add_paragraph()  # 空行
            return True
        except Exception as e:
            print(f"添加圖片失敗 {image_path}: {e}")
            # 添加圖片佔位符
            doc.add_paragraph(f"[圖片：{caption}]", style='Caption')
            return False
    else:
        # 添加圖片佔位符
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder_run = placeholder.add_run(f"[圖 {caption}：{os.path.basename(image_path)}]")
        placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
        placeholder_run.italic = True
        return False

def create_complete_quantum_word_document():
    """創建完整的量子廣義相對論 Word 文檔"""
    
    # 創建圖片目錄
    images_dir = create_images_directory()
    
    # 創建文檔
    doc = Document()
    
    # 設置中文字體
    style = doc.styles['Normal']
    style.font.name = 'PMingLiU'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
    style.font.size = Pt(12)
    
    # 設置頁面
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    print("開始生成完整 Word 文檔...")
    
    # 讀取三個章節文件
    chapter_files = [
        '量子廣義相對論_修改版_第1-2章.md',
        '量子廣義相對論_修改版_第3-5章.md', 
        '量子廣義相對論_修改版_第6-9章.md'
    ]
    
    # 處理每個章節
    for chapter_file in chapter_files:
        if not os.path.exists(chapter_file):
            print(f"警告：找不到 {chapter_file}")
            continue
            
        print(f"處理 {chapter_file}...")
        
        with open(chapter_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        for line_num, line in enumerate(lines):
            line = line.rstrip()
            
            if not line:
                continue
            
            # 處理標題
            if line.startswith('# ') and '量子廣義相對論' in line:
                # 主標題
                h = doc.add_heading(line[2:], 0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif line.startswith('## 第'):
                # 章標題
                doc.add_heading(line[3:], 1)
                
            elif line.startswith('### '):
                # 節標題
                doc.add_heading(line[4:], 2)
                
            elif line.startswith('#### '):
                # 小節標題
                doc.add_heading(line[5:], 3)
            
            # 作者行
            elif '作者：劉哲廷' in line:
                p = doc.add_paragraph(line.replace('**', ''))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
            
            # 圖片插入點
            elif '**圖 2：曲率分布圖**' in line or '圖 2：曲率分布圖' in line:
                add_image_to_document(doc, images_dir / 'figure_2_curvature.png', '2：曲率分布圖')
                
            elif '**圖 3：能量密度分布**' in line or '圖 3：能量密度分布' in line:
                add_image_to_document(doc, images_dir / 'figure_3_energy.png', '3：能量密度分布')
                
            elif '**圖 4：損失函數收斂曲線**' in line or '圖 4：損失函數收斂曲線' in line:
                add_image_to_document(doc, images_dir / 'figure_4_convergence.png', '4：損失函數收斂曲線')
                
            elif '**圖 5：AI 架構對映圖**' in line or '圖 5：模型層與理論層對映圖' in line:
                add_image_to_document(doc, images_dir / 'figure_5_mapping.png', '5：模型層與理論層對映圖')
                
            elif '**圖 6：量子修正項分布**' in line or '圖 6：量子修正項分布' in line:
                add_image_to_document(doc, images_dir / 'figure_6_quantum_correction.png', '6：量子修正項分布')
                
            elif '**圖 7：觀測者場效應**' in line or '圖 7：觀測者場效應' in line:
                add_image_to_document(doc, images_dir / 'figure_7_observer_field.png', '7：觀測者場效應')
                
            elif '**圖 8：理論層次結構圖**' in line or '圖 8：理論層次結構圖' in line:
                add_image_to_document(doc, images_dir / 'figure_8_hierarchy.png', '8：理論層次結構圖')
            
            # 分隔線
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            
            # 數學公式處理
            elif line.startswith('$$'):
                formula = line.replace('$$', '').strip()
                if formula:
                    # 轉換數學符號
                    converted_formula = convert_math_symbols(formula)
                    
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(converted_formula)
                    r.font.italic = True
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
            
            # 代碼塊（忽略）
            elif line.startswith('```'):
                continue
            
            # 表格行
            elif line.startswith('|') and '|' in line[1:]:
                # 簡單表格處理
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    table_text = ' | '.join(cells)
                    doc.add_paragraph(table_text, style='List Bullet')
            
            # 強調文字（紅色）
            elif '（本報告為個人理論探索' in line:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line)
                r.font.color.rgb = RGBColor(200, 0, 0)
                r.italic = True
                r.font.size = Pt(10)
            
            # 粗體標記處理
            elif '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.*?)\*\*', line)
                for i, part in enumerate(parts):
                    if part:
                        r = p.add_run(part)
                        if i % 2 == 1:  # 粗體
                            r.bold = True
            
            # 普通段落
            else:
                if not line.startswith('**圖') and not line.startswith('**('):
                    # 轉換數學符號
                    converted_line = convert_math_symbols(line)
                    doc.add_paragraph(converted_line)
    
    # 保存文檔
    output_file = '量子廣義相對論_完整版_含圖片_劉哲廷.docx'
    doc.save(output_file)
    
    print()
    print(f"✅ 完整 Word 文檔已生成：{output_file}")
    print()
    print("📋 後續步驟：")
    print("1. 將對應圖片放入 'images' 目錄")
    print("2. 用 Word 開啟文檔檢查格式")
    print("3. 手動調整數學公式格式（如果需要）")
    print("4. 檢查圖片顯示是否正常")
    
    return output_file

def create_image_instructions():
    """創建圖片說明文件"""
    instructions = """
# 圖片插入說明

## 需要的圖片文件

請將以下圖片放入 `images` 目錄：

### 1. figure_2_curvature.png
- **說明：** 曲率分布圖
- **內容：** 顯示熵場導出的時空彎曲結構
- **來源：** 第2章圖2

### 2. figure_3_energy.png  
- **說明：** 能量密度分布圖
- **內容：** 幾何與能量守恆關係
- **來源：** 第3章圖3

### 3. figure_4_convergence.png
- **說明：** 損失函數收斂曲線
- **內容：** 模型達到協變與能量平衡
- **來源：** 第4章圖4

### 4. figure_5_mapping.png
- **說明：** 模型層與理論層對映圖
- **內容：** AI層與物理層的結構映射
- **來源：** 第5章圖5

### 5. figure_6_quantum_correction.png
- **說明：** 量子修正項分布
- **內容：** 熵梯度導致的局域量子修正
- **來源：** 第3章圖6

### 6. figure_7_observer_field.png
- **說明：** 觀測者場效應圖
- **內容：** 觀測行為對時空的回饋作用
- **來源：** 第6章圖7

### 7. figure_8_hierarchy.png
- **說明：** 理論層次結構圖
- **內容：** 從QIGG到CSC的遞階演化
- **來源：** 第7章圖8

## 圖片要求

- **格式：** PNG 或 JPG
- **解析度：** 至少 300 DPI（用於列印）
- **尺寸：** 建議寬度 800-1200 像素
- **背景：** 白色或透明

## 替代方案

如果沒有圖片文件，文檔會自動插入佔位符：
`[圖 X：filename.png]`

可以在 Word 中手動插入對應圖片。

## 數學公式處理

文檔已自動轉換常見數學符號：
- `\\mu` → μ
- `\\nu` → ν  
- `\\rho` → ρ
- `\\nabla` → ∇
- `\\Omega` → Ω
- 等等...

如需更複雜的公式，建議使用 Word 的方程式編輯器。
"""
    
    with open("圖片插入說明.md", "w", encoding="utf-8") as f:
        f.write(instructions)
    
    print("✅ 圖片說明文件已創建：圖片插入說明.md")

if __name__ == '__main__':
    print("=" * 60)
    print("量子廣義相對論 - 完整 Word 文檔生成器")
    print("=" * 60)
    print()
    
    # 創建圖片說明
    create_image_instructions()
    print()
    
    # 生成完整文檔
    output_file = create_complete_quantum_word_document()
    
    print()
    print("🎉 完成！")
    print(f"📄 文檔位置：{output_file}")
    print("📁 圖片目錄：images/")
    print("📋 說明文件：圖片插入說明.md")
