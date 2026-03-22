#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Simple Word Document Generator
Basic version without complex regex
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    print("python-docx installed")
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def create_images_directory():
    """Create images directory"""
    images_dir = Path("images")
    images_dir.mkdir(exist_ok=True)
    
    print(f"Images directory created: {images_dir}")
    return images_dir

def simple_math_convert(text):
    """Simple math symbol conversion"""
    # Basic conversions only
    conversions = {
        '\\mu': 'μ',
        '\\nu': 'ν', 
        '\\rho': 'ρ',
        '\\lambda': 'λ',
        '\\kappa': 'κ',
        '\\Omega': 'Ω',
        '\\nabla': '∇',
        '\\Box': '□',
        '\\Psi': 'Ψ',
        '\\pi': 'π',
        '\\alpha': 'α',
        '\\beta': 'β',
        '\\gamma': 'γ',
        '\\delta': 'δ',
        '\\partial': '∂',
        '\\infty': '∞',
        '\\times': '×',
        '\\cdot': '·',
        '\\leq': '≤',
        '\\geq': '≥',
        '\\neq': '≠',
        '\\approx': '≈',
        '\\rightarrow': '→',
        '\\Rightarrow': '⇒',
        '\\langle': '⟨',
        '\\rangle': '⟩'
    }
    
    for latex, symbol in conversions.items():
        text = text.replace(latex, symbol)
    
    return text

def add_image_placeholder(doc, image_name, caption):
    """Add image placeholder"""
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_run = placeholder.add_run(f"[圖 {caption}：{image_name}]")
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    placeholder_run.italic = True
    placeholder_run.font.size = Pt(10)
    doc.add_paragraph()  # Empty line

def create_simple_word_document():
    """Create simple Word document"""
    
    # Create images directory
    images_dir = create_images_directory()
    
    # Create document
    doc = Document()
    
    # Set Chinese font
    style = doc.styles['Normal']
    style.font.name = 'PMingLiU'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
    style.font.size = Pt(12)
    
    print("Starting Word document generation...")
    
    # Read chapter files
    chapter_files = [
        '量子廣義相對論_修改版_第1-2章.md',
        '量子廣義相對論_修改版_第3-5章.md', 
        '量子廣義相對論_修改版_第6-9章.md'
    ]
    
    # Process each chapter
    for chapter_file in chapter_files:
        if not os.path.exists(chapter_file):
            print(f"Warning: {chapter_file} not found")
            continue
            
        print(f"Processing {chapter_file}...")
        
        with open(chapter_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        for line_num, line in enumerate(lines):
            line = line.rstrip()
            
            if not line:
                continue
            
            # Process headings
            if line.startswith('# ') and '量子廣義相對論' in line:
                h = doc.add_heading(line[2:], 0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif line.startswith('## 第'):
                doc.add_heading(line[3:], 1)
                
            elif line.startswith('### '):
                doc.add_heading(line[4:], 2)
                
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 3)
            
            # Author line
            elif '作者：劉哲廷' in line:
                p = doc.add_paragraph(line.replace('**', ''))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
            
            # Image insertion points
            elif '圖 2：曲率分布圖' in line:
                add_image_placeholder(doc, 'figure_2_curvature.png', '2：曲率分布圖')
                
            elif '圖 3：能量密度分布' in line:
                add_image_placeholder(doc, 'figure_3_energy.png', '3：能量密度分布')
                
            elif '圖 4：損失函數收斂曲線' in line:
                add_image_placeholder(doc, 'figure_4_convergence.png', '4：損失函數收斂曲線')
                
            elif '圖 5：' in line and '對映' in line:
                add_image_placeholder(doc, 'figure_5_mapping.png', '5：模型層與理論層對映圖')
                
            elif '圖 6：量子修正項分布' in line:
                add_image_placeholder(doc, 'figure_6_quantum_correction.png', '6：量子修正項分布')
                
            elif '圖 7：觀測者場效應' in line:
                add_image_placeholder(doc, 'figure_7_observer_field.png', '7：觀測者場效應')
                
            elif '圖 8：理論層次結構圖' in line:
                add_image_placeholder(doc, 'figure_8_hierarchy.png', '8：理論層次結構圖')
            
            # Separators
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            
            # Math formulas
            elif line.startswith('$$'):
                formula = line.replace('$$', '').strip()
                if formula:
                    # Simple math conversion
                    converted_formula = simple_math_convert(formula)
                    
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(converted_formula)
                    r.font.italic = True
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
            
            # Code blocks (ignore)
            elif line.startswith('```'):
                continue
            
            # Table rows (simplified)
            elif line.startswith('|') and '|' in line[1:]:
                # Just add as regular text
                clean_line = line.replace('|', ' | ')
                doc.add_paragraph(clean_line)
            
            # Emphasized text (red)
            elif '本報告為個人理論探索' in line:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line)
                r.font.color.rgb = RGBColor(200, 0, 0)
                r.italic = True
                r.font.size = Pt(10)
            
            # Bold text processing (simple)
            elif '**' in line:
                p = doc.add_paragraph()
                # Simple bold handling
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        r = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold
                            r.bold = True
            
            # Regular paragraphs
            else:
                if not line.startswith('**圖') and not line.startswith('**('):
                    # Simple math conversion
                    converted_line = simple_math_convert(line)
                    doc.add_paragraph(converted_line)
    
    # Save document
    output_file = 'Quantum_GR_Simple_Report.docx'
    doc.save(output_file)
    
    print()
    print(f"Word document generated: {output_file}")
    print()
    print("Next steps:")
    print("1. Place corresponding images in 'images' directory")
    print("2. Open document in Word")
    print("3. Replace placeholders with actual images")
    print("4. Adjust math formulas as needed")
    
    return output_file

def create_image_guide():
    """Create image guide"""
    guide = """
# Image Guide for Quantum GR Report

## Required Images

Place these images in the 'images' directory:

1. **figure_2_curvature.png** - Curvature Distribution Map
2. **figure_3_energy.png** - Energy Density Distribution  
3. **figure_4_convergence.png** - Loss Function Convergence Curve
4. **figure_5_mapping.png** - Model Layer Mapping
5. **figure_6_quantum_correction.png** - Quantum Correction Distribution
6. **figure_7_observer_field.png** - Observer Field Effect
7. **figure_8_hierarchy.png** - Theoretical Hierarchy

## How to Add Images

1. Save images as PNG files in the 'images' folder
2. Open the Word document
3. Find placeholders like [圖 2：曲率分布圖：figure_2_curvature.png]
4. Delete the placeholder text
5. Insert the actual image (Insert > Picture)
6. Center the image and add caption

## Math Formulas

The document includes basic math symbol conversion:
- μ, ν, ρ, λ, κ, Ω, ∇, π, etc.

For complex formulas, use Word's equation editor:
Insert > Equation

## Final Steps

1. Check all formatting
2. Verify image placement
3. Review math formulas
4. Add page numbers
5. Check Chinese font display
"""
    
    with open("Image_Guide.txt", "w", encoding="utf-8") as f:
        f.write(guide)
    
    print("Image guide created: Image_Guide.txt")

if __name__ == '__main__':
    print("=" * 60)
    print("Quantum GR - Simple Word Document Generator")
    print("=" * 60)
    print()
    
    # Create image guide
    create_image_guide()
    print()
    
    # Generate document
    output_file = create_simple_word_document()
    
    print()
    print("Complete!")
    print(f"Document: {output_file}")
    print("Images directory: images/")
    print("Guide: Image_Guide.txt")
