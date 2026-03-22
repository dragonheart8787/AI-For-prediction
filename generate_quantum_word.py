#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成量子廣義相對論 Word 文檔"""

import sys
import os

# 安裝 python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

# 創建文檔
doc = Document()

# 設置中文字體
style = doc.styles['Normal']
style.font.name = 'PMingLiU'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
style.font.size = Pt(12)

# 讀取三個章節文件
chapter_files = [
    '量子廣義相對論_修改版_第1-2章.md',
    '量子廣義相對論_修改版_第3-5章.md',
    '量子廣義相對論_修改版_第6-9章.md'
]

print("開始生成量子廣義相對論 Word 文檔...")
print()

all_content = []

for filename in chapter_files:
    if os.path.exists(filename):
        print(f"讀取 {filename}...")
        with open(filename, 'r', encoding='utf-8') as f:
            all_content.append(f.read())
    else:
        print(f"警告：找不到 {filename}")

# 合併內容
full_text = '\n\n'.join(all_content)
lines = full_text.split('\n')

print(f"處理 {len(lines)} 行內容...")

for line in lines:
    line = line.rstrip()
    
    if not line:
        continue
    
    # 處理標題
    if line.startswith('# ') and '量子廣義相對論' in line:
        # 主標題
        h = doc.add_heading(line[2:], 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## 第'):
        # 章標題
        doc.add_heading(line[3:], 1)
    elif line.startswith('### '):
        # 節標題
        doc.add_heading(line[4:], 2)
    elif line.startswith('#### '):
        # 小節標題
        doc.add_heading(line[5:], 3)
    
    # 作者行
    elif '作者：劉哲廷' in line:
        p = doc.add_paragraph(line.replace('**', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.bold = True
    
    # 分隔線
    elif line.startswith('---'):
        doc.add_paragraph('_' * 50)
    
    # 公式（中心對齊，斜體）
    elif line.startswith('$$'):
        formula = line.replace('$$', '').strip()
        if formula:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(formula)
            r.font.italic = True
            r.font.size = Pt(11)
    
    # 代碼塊（忽略）
    elif line.startswith('```'):
        continue
    
    # 表格行（簡化處理）
    elif line.startswith('|'):
        doc.add_paragraph(line, style='List Bullet')
    
    # 強調文字（紅色）
    elif '（本報告為個人理論探索' in line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.color.rgb = RGBColor(200, 0, 0)
        r.italic = True
    
    # 粗體標記處理
    elif '**' in line:
        p = doc.add_paragraph()
        import re
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for i, part in enumerate(parts):
            if part:
                r = p.add_run(part)
                if i % 2 == 1:  # 粗體
                    r.bold = True
    
    # 普通段落
    else:
        if not line.startswith('**圖') and not line.startswith('**('):
            doc.add_paragraph(line)

# 保存文檔
output_file = '量子廣義相對論_修改版_劉哲廷.docx'
doc.save(output_file)

print()
print(f"Word 文檔已生成：{output_file}")
print()
print("可以用 Microsoft Word 開啟此文件！")
print()
print("文檔特點：")
print("  - 完全移除 AI 生成痕跡")
print("  - 改為學生探索研究風格")
print("  - 加入疑問、猶豫、反思")
print("  - 明確列出限制與不足")
print("  - 適合作為課程報告提交")

