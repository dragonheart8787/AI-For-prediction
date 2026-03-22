#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Complete Quantum GR Word Document Generator
Includes image insertion and math symbol conversion
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    print("python-docx installed")
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def create_images_directory():
    """Create images directory"""
    images_dir = Path("images")
    images_dir.mkdir(exist_ok=True)
    
    # Create image descriptions file
    image_descriptions = {
        "figure_2_curvature.png": "曲率分布圖 - 顯示熵場導出的時空彎曲結構",
        "figure_3_energy.png": "能量密度分布 - 幾何與能量守恆關係", 
        "figure_4_convergence.png": "損失函數收斂曲線 - 模型達到協變與能量平衡",
        "figure_5_mapping.png": "模型層與理論層對映圖 - AI層與物理層的結構映射",
        "figure_6_quantum_correction.png": "量子修正項分布 - 熵梯度導致的局域量子修正",
        "figure_7_observer_field.png": "觀測者場效應 - 觀測行為對時空的回饋作用",
        "figure_8_hierarchy.png": "理論層次結構圖 - 從QIGG到CSC的遞階演化"
    }
    
    # Create image descriptions file
    with open(images_dir / "image_descriptions.txt", "w", encoding="utf-8") as f:
        f.write("Image Descriptions\n")
        f.write("=" * 50 + "\n\n")
        for filename, description in image_descriptions.items():
            f.write(f"{filename}: {description}\n")
    
    print(f"Images directory created: {images_dir}")
    return images_dir

def convert_math_symbols(text):
    """Convert LaTeX math symbols to readable format"""
    # Common math symbol conversions
    math_conversions = {
        # Greek letters
        r'\\mu': 'μ',
        r'\\nu': 'ν', 
        r'\\rho': 'ρ',
        r'\\lambda': 'λ',
        r'\\kappa': 'κ',
        r'\\Omega': 'Ω',
        r'\\nabla': '∇',
        r'\\Box': '□',
        r'\\Psi': 'Ψ',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\gamma': 'γ',
        r'\\delta': 'δ',
        r'\\epsilon': 'ε',
        r'\\pi': 'π',
        r'\\sigma': 'σ',
        r'\\phi': 'φ',
        r'\\psi': 'ψ',
        r'\\omega': 'ω',
        
        # Math operators
        r'\\partial': '∂',
        r'\\sum': '∑',
        r'\\int': '∫',
        r'\\infty': '∞',
        r'\\times': '×',
        r'\\cdot': '·',
        r'\\pm': '±',
        r'\\leq': '≤',
        r'\\geq': '≥',
        r'\\neq': '≠',
        r'\\approx': '≈',
        r'\\equiv': '≡',
        r'\\propto': '∝',
        r'\\in': '∈',
        r'\\rightarrow': '→',
        r'\\leftarrow': '←',
        r'\\Rightarrow': '⇒',
        r'\\Leftrightarrow': '⇔',
        
        # Brackets
        r'\\left\\(': '(',
        r'\\right\\)': ')',
        r'\\left\\[': '[',
        r'\\right\\]': ']',
        r'\\left\\{': '{',
        r'\\right\\}': '}',
        
        # Fractions
        r'\\frac\{([^}]+)\}\{([^}]+)\}': r'\1/\2',
        
        # Subscripts/superscripts
        r'\^([0-9])': r'^\1',
        r'_([a-zA-Z0-9]+)': r'_\1',
        
        # Other
        r'\\langle': '⟨',
        r'\\rangle': '⟩',
        r'\\mid': '|'
    }
    
    for pattern, replacement in math_conversions.items():
        text = re.sub(pattern, replacement, text)
    
    return text

def add_image_to_document(doc, image_path, caption, width=4.5):
    """Add image to document"""
    if os.path.exists(image_path):
        try:
            # Add image
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(image_path, width=Inches(width))
            
            # Add caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(f"圖 {caption}")
            caption_run.font.size = Pt(11)
            caption_run.italic = True
            
            doc.add_paragraph()  # Empty line
            return True
        except Exception as e:
            print(f"Failed to add image {image_path}: {e}")
            # Add placeholder
            doc.add_paragraph(f"[圖片：{caption}]", style='Caption')
            return False
    else:
        # Add placeholder
        placeholder = doc.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder_run = placeholder.add_run(f"[圖 {caption}：{os.path.basename(image_path)}]")
        placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
        placeholder_run.italic = True
        return False

def create_complete_quantum_word_document():
    """Create complete quantum GR Word document"""
    
    # Create images directory
    images_dir = create_images_directory()
    
    # Create document
    doc = Document()
    
    # Set Chinese font
    style = doc.styles['Normal']
    style.font.name = 'PMingLiU'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
    style.font.size = Pt(12)
    
    print("Starting complete Word document generation...")
    
    # Read chapter files
    chapter_files = [
        '量子廣義相對論_修改版_第1-2章.md',
        '量子廣義相對論_修改版_第3-5章.md', 
        '量子廣義相對論_修改版_第6-9章.md'
    ]
    
    # Process each chapter
    for chapter_file in chapter_files:
        if not os.path.exists(chapter_file):
            print(f"Warning: {chapter_file} not found")
            continue
            
        print(f"Processing {chapter_file}...")
        
        with open(chapter_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        for line_num, line in enumerate(lines):
            line = line.rstrip()
            
            if not line:
                continue
            
            # Process headings
            if line.startswith('# ') and '量子廣義相對論' in line:
                h = doc.add_heading(line[2:], 0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            elif line.startswith('## 第'):
                doc.add_heading(line[3:], 1)
                
            elif line.startswith('### '):
                doc.add_heading(line[4:], 2)
                
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 3)
            
            # Author line
            elif '作者：劉哲廷' in line:
                p = doc.add_paragraph(line.replace('**', ''))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
            
            # Image insertion points
            elif '**圖 2：曲率分布圖**' in line or '圖 2：曲率分布圖' in line:
                add_image_to_document(doc, images_dir / 'figure_2_curvature.png', '2：曲率分布圖')
                
            elif '**圖 3：能量密度分布**' in line or '圖 3：能量密度分布' in line:
                add_image_to_document(doc, images_dir / 'figure_3_energy.png', '3：能量密度分布')
                
            elif '**圖 4：損失函數收斂曲線**' in line or '圖 4：損失函數收斂曲線' in line:
                add_image_to_document(doc, images_dir / 'figure_4_convergence.png', '4：損失函數收斂曲線')
                
            elif '**圖 5：AI 架構對映圖**' in line or '圖 5：模型層與理論層對映圖' in line:
                add_image_to_document(doc, images_dir / 'figure_5_mapping.png', '5：模型層與理論層對映圖')
                
            elif '**圖 6：量子修正項分布**' in line or '圖 6：量子修正項分布' in line:
                add_image_to_document(doc, images_dir / 'figure_6_quantum_correction.png', '6：量子修正項分布')
                
            elif '**圖 7：觀測者場效應**' in line or '圖 7：觀測者場效應' in line:
                add_image_to_document(doc, images_dir / 'figure_7_observer_field.png', '7：觀測者場效應')
                
            elif '**圖 8：理論層次結構圖**' in line or '圖 8：理論層次結構圖' in line:
                add_image_to_document(doc, images_dir / 'figure_8_hierarchy.png', '8：理論層次結構圖')
            
            # Separators
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            
            # Math formulas
            elif line.startswith('$$'):
                formula = line.replace('$$', '').strip()
                if formula:
                    # Convert math symbols
                    converted_formula = convert_math_symbols(formula)
                    
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(converted_formula)
                    r.font.italic = True
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
            
            # Code blocks (ignore)
            elif line.startswith('```'):
                continue
            
            # Table rows
            elif line.startswith('|') and '|' in line[1:]:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    table_text = ' | '.join(cells)
                    doc.add_paragraph(table_text, style='List Bullet')
            
            # Emphasized text (red)
            elif '（本報告為個人理論探索' in line:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line)
                r.font.color.rgb = RGBColor(200, 0, 0)
                r.italic = True
                r.font.size = Pt(10)
            
            # Bold text processing
            elif '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.*?)\*\*', line)
                for i, part in enumerate(parts):
                    if part:
                        r = p.add_run(part)
                        if i % 2 == 1:  # Bold
                            r.bold = True
            
            # Regular paragraphs
            else:
                if not line.startswith('**圖') and not line.startswith('**('):
                    # Convert math symbols
                    converted_line = convert_math_symbols(line)
                    doc.add_paragraph(converted_line)
    
    # Save document
    output_file = 'Quantum_GR_Complete_Report.docx'
    doc.save(output_file)
    
    print()
    print(f"Complete Word document generated: {output_file}")
    print()
    print("📋 Next steps:")
    print("1. Place corresponding images in 'images' directory")
    print("2. Open document in Word to check formatting")
    print("3. Manually adjust math formula formatting if needed")
    print("4. Check if images display correctly")
    
    return output_file

def create_image_instructions():
    """Create image instructions file"""
    instructions = """
# Image Insertion Instructions

## Required Image Files

Please place the following images in the `images` directory:

### 1. figure_2_curvature.png
- **Description:** Curvature Distribution Map
- **Content:** Shows spacetime curved structure derived from entropy field
- **Source:** Chapter 2, Figure 2

### 2. figure_3_energy.png  
- **Description:** Energy Density Distribution
- **Content:** Relationship between geometry and energy conservation
- **Source:** Chapter 3, Figure 3

### 3. figure_4_convergence.png
- **Description:** Loss Function Convergence Curve
- **Content:** Model achieves covariance and energy balance
- **Source:** Chapter 4, Figure 4

### 4. figure_5_mapping.png
- **Description:** Model Layer and Theoretical Layer Mapping
- **Content:** Structural mapping between AI layer and physical layer
- **Source:** Chapter 5, Figure 5

### 5. figure_6_quantum_correction.png
- **Description:** Quantum Correction Term Distribution
- **Content:** Local quantum correction caused by entropy gradient
- **Source:** Chapter 3, Figure 6

### 6. figure_7_observer_field.png
- **Description:** Observer Field Effect
- **Content:** Feedback effect of observational behavior on spacetime
- **Source:** Chapter 6, Figure 7

### 7. figure_8_hierarchy.png
- **Description:** Theoretical Hierarchical Structure
- **Content:** Hierarchical evolution from QIGG to CSC
- **Source:** Chapter 7, Figure 8

## Image Requirements

- **Format:** PNG or JPG
- **Resolution:** At least 300 DPI (for printing)
- **Size:** Recommended width 800-1200 pixels
- **Background:** White or transparent

## Alternative

If no image files are available, the document will automatically insert placeholders:
`[圖 X：filename.png]`

You can manually insert corresponding images in Word.

## Math Formula Processing

The document automatically converts common math symbols:
- `\\mu` → μ
- `\\nu` → ν  
- `\\rho` → ρ
- `\\nabla` → ∇
- `\\Omega` → Ω
- etc.

For more complex formulas, it's recommended to use Word's equation editor.
"""
    
    with open("Image_Instructions.md", "w", encoding="utf-8") as f:
        f.write(instructions)
    
    print("Image instructions file created: Image_Instructions.md")

if __name__ == '__main__':
    print("=" * 60)
    print("Quantum GR - Complete Word Document Generator")
    print("=" * 60)
    print()
    
    # Create image instructions
    create_image_instructions()
    print()
    
    # Generate complete document
    output_file = create_complete_quantum_word_document()
    
    print()
    print("Complete!")
    print(f"📄 Document location: {output_file}")
    print("📁 Images directory: images/")
    print("📋 Instructions file: Image_Instructions.md")
