#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys

print("Word Document Generator")
print("=" * 60)
print()

# Check if python-docx is installed
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    print("python-docx installed")
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

# Create document
doc = Document()

# Set Chinese font
style = doc.styles['Normal']
style.font.name = 'PMingLiU'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')
style.font.size = Pt(12)

# Read all markdown files
files = [
    'Complete_Academic_Report_Part1.md',
    'Complete_Academic_Report_Part2.md',
    'Complete_Academic_Report_Part3.md',
    'Complete_Academic_Report_Part4.md',
    'Complete_Academic_Report_Part5.md'
]

quantum_files = [
    'quantum_report_ch1-2.md',
    'quantum_report_ch3-5.md',
    'quantum_report_ch6-9.md'
]

import os
if os.path.exists(quantum_files[0]):
    files = quantum_files
    output_name = 'Quantum_GR_Report.docx'
else:
    output_name = 'SuperFusionAGI_Report.docx'

print(f"Processing {len(files)} files...")

for filename in files:
    if not os.path.exists(filename):
        print(f"Skipping {filename} (not found)")
        continue
        
    print(f"Reading {filename}...")
    
    with open(filename, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 3)
        # Code blocks or formulas
        elif line.startswith('```') or line.startswith('$$'):
            continue
        # Horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        # Tables (simplified)
        elif line.startswith('|'):
            doc.add_paragraph(line)
        # Normal paragraphs
        else:
            if not line.startswith('**Fig') and not line.startswith('**圖'):
                doc.add_paragraph(line)

# Save
doc.save(output_name)
print()
print(f"Word document created: {output_name}")
print()
print("You can now open this file with Microsoft Word!")

